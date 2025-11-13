#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
KPMG Workbench 戦略評価フレームワーク Word文書生成器（日本語版）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """段落にハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # ハイパーリンクスタイル（青色下線）を設定
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

    return hyperlink

def set_heading_style(doc):
    """見出しスタイルを設定"""
    styles = doc.styles

    # 見出し1スタイル
    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    # 東アジアフォント設定
    rPr = h1._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        h1._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:eastAsia'), 'Meiryo UI')

    # 見出し2スタイル
    h2 = styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 112, 192)
    # 東アジアフォント設定
    rPr = h2._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        h2._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:eastAsia'), 'Meiryo UI')

    # 見出し3スタイル
    h3 = styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(12)
    h3.font.bold = True
    # 東アジアフォント設定
    rPr = h3._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        h3._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:eastAsia'), 'Meiryo UI')

def create_cover_page(doc):
    """表紙ページを作成"""
    # タイトル
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('KPMG Workbench 戦略評価レポート')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # 空行

    # サブタイトル
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI開発プラットフォーム詳細評価フレームワーク')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # メタ情報
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('評価担当者: [氏名]\n').font.size = Pt(12)
    info.add_run('役職: Senior Consultant, AI Development\n').font.size = Pt(12)
    info.add_run('日付: [評価日]\n').font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # バージョン情報
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

def create_toc(doc):
    """目次ページを作成"""
    doc.add_heading('目次', 0)

    toc = doc.add_paragraph()
    toc.add_run('【Wordで目次を生成する手順】\n').font.italic = True
    toc.add_run('1. カーソルをここに置く\n').font.italic = True
    toc.add_run('2. 「参考資料」タブ → 「目次」 → 自動目次スタイルを選択\n').font.italic = True
    toc.add_run('3. 内容記入完了後、目次を右クリック → 「フィールド更新」 → 「目次をすべて更新」\n\n').font.italic = True
    toc.add_run('注: すべての章見出しは見出しスタイルに設定されており、自動的にハイパーリンク目次が生成されます').font.italic = True

    doc.add_page_break()

def add_dimension_0(doc):
    """維度0: 事前準備（日本語版・詳細）"""
    doc.add_heading('0. 事前準備: 学習パスと認証コスト評価', 1)

    doc.add_paragraph('【評価背景】これはチームメンバーにとって最大の参入障壁であり、時間コストと学習品質を詳細に評価する必要があります。')

    doc.add_heading('0.1 認証要件', 2)
    p = doc.add_paragraph('KPMG Workbenchにアクセスするには、以下の認証パスを完了する必要があります (出典: ')
    add_hyperlink(p, 'KPMG Workbench Learning & Development Hub',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-development.aspx')
    p.add_run(')：')

    doc.add_paragraph()
    doc.add_heading('必須要件', 3)

    doc.add_paragraph('✓ 以下のいずれかの学習パスを完了し、KPMG Workbench Knowledge Badgeを取得すること：', style='List Bullet')

    p1 = doc.add_paragraph('  • ', style='List Bullet 2')
    add_hyperlink(p1, 'Developer Learning Path',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-development-track.aspx')
    p1.add_run(' （開発者向け）')

    p2 = doc.add_paragraph('  • ', style='List Bullet 2')
    add_hyperlink(p2, 'Product Management Learning Path',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-product-management-track.aspx')
    p2.add_run(' （プロダクトマネージャー向け）')

    doc.add_paragraph()
    doc.add_heading('事前トレーニング要件', 3)
    doc.add_paragraph('• 実務経験のあるエンジニア、技術者、またはプロダクトスペシャリストであること', style='List Bullet')

    p3 = doc.add_paragraph('• ', style='List Bullet')
    add_hyperlink(p3, 'GitHub EMU',
                  'https://handbook.code.kpmg.com/KPMG-Code/GitHub/Organization%20onboarding/')
    p3.add_run(' にオンボーディングされていること')

    doc.add_paragraph('• GitHub EMUリポジトリに少なくとも1つのPull Requestを提出していること', style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('推奨認証（必須ではない）', 3)
    doc.add_paragraph('Knowledge Badge トレーニングを開始する前に、以下の認証のうち2つ以上を完了することが推奨されます：')

    doc.add_heading('開発者向け推奨認証:', 4)
    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Azure Fundamentals AZ-900',
                  'https://learn.microsoft.com/en-us/credentials/certifications/azure-fundamentals/?practice-assessment-type=certification')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Azure AI Fundamentals AI-900',
                  'https://learn.microsoft.com/en-us/credentials/certifications/azure-ai-fundamentals/?practice-assessment-type=certification')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'GitHub Foundations',
                  'https://learn.microsoft.com/en-us/collections/o1njfe825p602p')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'GitHub Actions',
                  'https://learn.microsoft.com/en-us/collections/n5p4a5z7keznp5')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Responsible AI',
                  'https://app.pluralsight.com/library/courses/artificial-intelligence-essentials-responsible-ai/table-of-contents')

    doc.add_heading('プロダクトマネージャー向け推奨認証:', 4)
    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Azure Fundamentals AZ-900',
                  'https://learn.microsoft.com/en-us/credentials/certifications/azure-fundamentals/?practice-assessment-type=certification')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Professional Scrum Master PSM I',
                  'https://www.scrum.org/assessments/professional-scrum-master-i-certification')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Professional Scrum Product Owner PSPO I',
                  'https://www.scrum.org/assessments/professional-scrum-product-owner-i-certification')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'GitHub Foundations',
                  'https://learn.microsoft.com/en-us/collections/o1njfe825p602p')

    p = doc.add_paragraph('• ')
    add_hyperlink(p, 'Responsible AI',
                  'https://app.pluralsight.com/library/courses/artificial-intelligence-essentials-responsible-ai/table-of-contents')

    doc.add_page_break()

    # Developer Learning Path詳細
    doc.add_heading('0.2 Developer Learning Path 詳細', 2)
    p = doc.add_paragraph('プログラム名: GX25_PRO_KPMG Workbench for Developers\n')
    p = doc.add_paragraph('プログラムID: ')
    add_hyperlink(p, 'GX25_CFS_DDF_AI_BLDG_WB_D_PRO',
                  'https://kpmgic.lms.hr.cloud.sap/learning/user/learning/program/viewProgramDetails.do?fromSF=Y&programID=GX25_CFS_DDF_AI_BLDG_WB_D_PRO')
    p.add_run('\n総所要時間: 約5.3時間（318分）')

    doc.add_paragraph()
    doc.add_heading('モジュール一覧:', 3)

    # Developer modules with links
    dev_modules = [
        ('1. Introduction to KPMG Workbench', '54分'),
        ('2. Revolutionizing AI Productivity: Dive into KPMG Workbench', '35分'),
        ('3. Deep Dive: Inference API', '26分'),
        ('4. Deep Dive: Completion API', '28分'),
        ('5. RAG: Overview and Building Blocks', '49分'),
        ('6. RAG: Leading Practices', '53分'),
        ('7. Tailoring KPMG Workbench for Global: Feature Flags', '13分'),
        ('8. Designing AI Experiences with KPMG Workbench', '39分'),
        ('9. Building Better, Faster: Guide to Developer Resources', '21分')
    ]

    for i, (title, duration) in enumerate(dev_modules, 1):
        doc.add_paragraph(f'{title} - {duration}', style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('重要な注意事項:', 3)
    p = doc.add_paragraph()
    p.add_run('ビデオを完全に視聴してください（95%）。完了ステータスは').font.bold = False
    p.add_run('24～48時間後').font.bold = True
    p.add_run('にシステムに転送されます。').font.bold = False

    p = doc.add_paragraph()
    p.add_run('原文: "').font.italic = True
    p.add_run('Please watch the video in full (95%) to ensure the completion is captured. Completions will be transferred after 24 - 48 hours.').font.italic = True
    p.add_run('"').font.italic = True

    doc.add_page_break()

    # Product Management Learning Path詳細
    doc.add_heading('0.3 Product Management Learning Path 詳細', 2)
    p = doc.add_paragraph('プログラム名: GX25_PRO_KPMG Workbench for Product Managers\n')
    p = doc.add_paragraph('プログラムID: ')
    add_hyperlink(p, 'GX25_CFS_DDF_AI_BLDG_WB_PM_PRO',
                  'https://kpmgic.lms.hr.cloud.sap/learning/user/learning/program/viewProgramDetails.do?fromSF=Y&programID=GX25_CFS_DDF_AI_BLDG_WB_PM_PRO')
    p.add_run('\n総所要時間: 約5.0時間（301分）')

    doc.add_paragraph()
    doc.add_heading('モジュール一覧:', 3)

    # PM modules
    pm_modules = [
        ('1. Introduction to KPMG Workbench', '54分'),
        ('2. Panel discussion', '45分'),
        ('3. Revolutionizing AI Productivity: Dive into KPMG Workbench', '35分'),
        ('4. Why Choose KPMG Workbench? Advancing your AI Innovations', '30分'),
        ('5. Safeguarding Innovation: IP and Patenting Strategies', '41分'),
        ('6. Microsoft Keynote - Agentic AI Thinking', '47分'),
        ('7. Migration Strategies: Transitioning to KPMG Workbench', '23分'),
        ('8. Submitting Feature Requests and Collaborating', '15分'),
        ('9. Support and Maintenance for Applications', '11分')
    ]

    for i, (title, duration) in enumerate(pm_modules, 1):
        doc.add_paragraph(f'{title} - {duration}', style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('重要な注意事項:', 3)
    p = doc.add_paragraph()
    p.add_run('ビデオを完全に視聴してください（95%）。完了ステータスは').font.bold = False
    p.add_run('24～48時間後').font.bold = True
    p.add_run('にシステムに転送されます。').font.bold = False

    p = doc.add_paragraph()
    p.add_run('原文: "').font.italic = True
    p.add_run('Please watch the video in full (95%) to ensure the completion is captured. Completions will be transferred after 24 - 48 hours.').font.italic = True
    p.add_run('"').font.italic = True

    doc.add_page_break()

    # 時間コスト評価
    doc.add_heading('0.4 時間コスト評価', 2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '学習モジュール'
    headers[1].text = '予定時間'
    headers[2].text = '実際の時間'

    modules = [
        ('事前要件（GitHub EMU等）', '___時間', ''),
        ('Developer / PM Learning Path', '5～5.3時間', ''),
        ('Assessment/Badge', '___時間', ''),
        ('合計', '___時間', '')
    ]

    for i, (module, est, actual) in enumerate(modules, 1):
        row = table.rows[i].cells
        row[0].text = module
        row[1].text = est
        row[2].text = actual

    doc.add_paragraph()

    doc.add_heading('0.5 API キー取得', 2)
    p = doc.add_paragraph('Badge取得後、')
    add_hyperlink(p, 'KPMG Workbench developer onboarding request form',
                  'https://kpmggoprod.service-now.com/sp?id=sc_cat_item&sys_id=623c6518c314a61088532485e0013117&sysparm_category=3cae446893230a10324c76847aba1033')
    p.add_run(' からAPIキーをリクエストします。')

    doc.add_paragraph('\n✓ Badgeの証明書を添付してください')
    doc.add_paragraph('✓ メンバーファーム承認者からの承認メールを添付してください')
    doc.add_paragraph('\nAPIキーとDeveloper Portalへのアクセスは、2～3営業日以内にメールで届きます。')

    doc.add_page_break()

def main():
    """メイン関数"""
    import sys
    import io

    # Windows コンソールエンコーディング問題を修正
    if sys.platform == 'win32':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    print('KPMG Workbench戦略評価レポート（日本語版）の生成を開始します...')

    # ドキュメントを作成
    doc = Document()

    # デフォルトフォントを設定
    doc.styles['Normal'].font.name = 'Arial'  # 英語
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo UI')  # 日本語
    doc.styles['Normal'].font.size = Pt(11)

    # 見出しスタイルを設定
    set_heading_style(doc)

    # 各セクションを作成
    print('  - 表紙ページを作成中...')
    create_cover_page(doc)

    print('  - 目次ページを作成中...')
    create_toc(doc)

    print('  - 維度0: 事前準備（詳細版）を追加中...')
    add_dimension_0(doc)

    # 保存
    output_path = r'C:\Users\junchenma\workbench-research\KPMG_Workbench戦略評価レポート_日本語版.docx'
    doc.save(output_path)

    print(f'✅ ドキュメント生成成功!')
    print(f'📁 ファイル場所: {output_path}')
    print('\n📝 次のステップ:')
    print('1. Wordでドキュメントを開く')
    print('2. カーソルを「目次」ページのプレースホルダーに置く')
    print('3. 参考資料 → 目次 → 自動目次スタイルを選択')
    print('4. 内容記入完了後、目次を右クリック → フィールド更新 → 目次をすべて更新')
    print('5. WordのPDF保存機能で最終版を出力')

if __name__ == '__main__':
    main()
