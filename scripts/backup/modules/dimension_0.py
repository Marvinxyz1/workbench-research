# -*- coding: utf-8 -*-
"""
KPMG Workbench 戦略評価フレームワーク - 維度0: 事前準備
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from utils import add_hyperlink
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# KPMG品牌颜色
KPMG_BLUE = RGBColor(0, 94, 184)  # KPMG品牌蓝色
KPMG_DARK_BLUE = RGBColor(0, 51, 141)  # KPMG深蓝色


def set_run_font(run, font_name='Arial', east_asia_font='Meiryo UI'):
    """Run要素に日本語フォントを設定

    Args:
        run: python-docx Run要素
        font_name: 欧文フォント名
        east_asia_font: 東アジアフォント名
    """
    run.font.name = font_name
    # 東アジアフォント設定
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), east_asia_font)


def set_table_style(table, has_header=True):
    """KPMGスタイルのテーブル設定

    Args:
        table: python-docx テーブルオブジェクト
        has_header: ヘッダー行があるかどうか（デフォルト: True）
    """
    # テーブルスタイルを適用
    table.style = 'Light Grid Accent 1'

    # ヘッダー行にKPMG青色の背景を設定
    if has_header and len(table.rows) > 0:
        header_cells = table.rows[0].cells
        for cell in header_cells:
            # 背景色を設定
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '005EB8')  # KPMG Blue
            cell._element.get_or_add_tcPr().append(shading_elm)

            # ヘッダーテキストを白色・太字に設定
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    set_run_font(run)

    # すべてのデータ行の太字を解除（第一列の自動太字を防ぐ）
    start_row = 1 if has_header else 0
    for row_idx in range(start_row, len(table.rows)):
        for cell in table.rows[row_idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = False
                    set_run_font(run)


def add_heading_with_font(doc, text, level):
    """日本語フォント付き見出しを追加（KPMG青色スタイル）

    Args:
        doc: python-docx文書オブジェクト
        text: 見出しテキスト
        level: 見出しレベル (0-3)

    Returns:
        作成された段落オブジェクト
    """
    heading = doc.add_heading(text, level)
    # 見出しの全てのRunにフォントとKPMG青色を設定
    for run in heading.runs:
        set_run_font(run)
        if level <= 2:  # レベル1と2の見出しにKPMG青色を適用
            run.font.color.rgb = KPMG_BLUE
    return heading


def add_dimension_0(doc):
    """維度0: 事前準備（日本語版・詳細）

    Args:
        doc: python-docx文書オブジェクト
    """
    add_heading_with_font(doc, '0. 事前準備: 学習パスと認証コスト評価', 1)

    doc.add_paragraph('KPMG WorkbenchにおけるAPI申請プロセスは、そのロジックが複雑であり、関連ページ間の連携も不明瞭であるため、以下の達成を目指します。')
    doc.add_paragraph()
    doc.add_paragraph('API申請の具体的な手順とプロセスを明確に抽出し、体系化する。')
    doc.add_paragraph()
    doc.add_paragraph('これにより、チームメンバーのAPI申請にかかる時間と学習コストを削減し、申請の迅速化を図る。')

    add_heading_with_font(doc, '0.1 認証要件', 2)
    p = doc.add_paragraph('KPMG Workbenchにアクセスするには、以下の認証パスを完了する必要があります (出典: ')
    add_hyperlink(p, 'KPMG Workbench Learning & Development Hub',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-development.aspx')
    p.add_run(')：')

    doc.add_paragraph()
    add_heading_with_font(doc, '必須要件', 3)

    doc.add_paragraph('以下のいずれかの学習パスを完了し、KPMG Workbench Knowledge Badgeを取得すること：', style='List Bullet')

    p1 = doc.add_paragraph(style='List Bullet 2')
    add_hyperlink(p1, 'Developer Learning Path',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-development-track.aspx')
    p1.add_run(' （開発者向け）')

    p2 = doc.add_paragraph(style='List Bullet 2')
    add_hyperlink(p2, 'Product Management Learning Path',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-product-management-track.aspx')
    p2.add_run(' （プロダクトマネージャー向け）')

    doc.add_paragraph()
    add_heading_with_font(doc, '事前トレーニング要件', 3)
    doc.add_paragraph('実務経験のあるエンジニア、技術者、またはプロダクトスペシャリストであること', style='List Bullet')

    p3 = doc.add_paragraph(style='List Bullet')
    add_hyperlink(p3, 'GitHub EMU',
                  'https://handbook.code.kpmg.com/KPMG-Code/GitHub/Organization%20onboarding/')
    p3.add_run(' にオンボーディングされていること')

    doc.add_paragraph('GitHub EMUリポジトリに少なくとも1つのPull Requestを提出していること', style='List Bullet')

    doc.add_paragraph()
    add_heading_with_font(doc, '推奨認証（必須ではない）', 3)
    doc.add_paragraph('Knowledge Badge トレーニングを開始する前に、以下の認証のうち2つ以上を完了することが推奨されます：')

    # 統合認証テーブル（2列）
    cert_table = doc.add_table(rows=6, cols=2)

    # ヘッダー行
    headers = cert_table.rows[0].cells
    headers[0].text = '開発者向け推奨認証'
    headers[1].text = 'プロダクトマネージャー向け推奨認証'

    # 開発者向け認証リスト
    dev_certs = [
        ('Azure Fundamentals AZ-900', 'https://learn.microsoft.com/en-us/credentials/certifications/azure-fundamentals/?practice-assessment-type=certification'),
        ('Azure AI Fundamentals AI-900', 'https://learn.microsoft.com/en-us/credentials/certifications/azure-ai-fundamentals/?practice-assessment-type=certification'),
        ('GitHub Foundations', 'https://learn.microsoft.com/en-us/collections/o1njfe825p602p'),
        ('GitHub Actions', 'https://learn.microsoft.com/en-us/collections/n5p4a5z7keznp5'),
        ('Responsible AI', 'https://app.pluralsight.com/library/courses/artificial-intelligence-essentials-responsible-ai/table-of-contents')
    ]

    # PM向け認証リスト
    pm_certs = [
        ('Azure Fundamentals AZ-900', 'https://learn.microsoft.com/en-us/credentials/certifications/azure-fundamentals/?practice-assessment-type=certification'),
        ('Professional Scrum Master PSM I', 'https://www.scrum.org/assessments/professional-scrum-master-i-certification'),
        ('Professional Scrum Product Owner PSPO I', 'https://www.scrum.org/assessments/professional-scrum-product-owner-i-certification'),
        ('GitHub Foundations', 'https://learn.microsoft.com/en-us/collections/o1njfe825p602p'),
        ('Responsible AI', 'https://app.pluralsight.com/library/courses/artificial-intelligence-essentials-responsible-ai/table-of-contents')
    ]

    # データ行を埋める
    for i in range(5):
        row = cert_table.rows[i + 1]

        # 開発者向け列
        if i < len(dev_certs):
            p_dev = row.cells[0].paragraphs[0]
            add_hyperlink(p_dev, dev_certs[i][0], dev_certs[i][1])

        # PM向け列
        if i < len(pm_certs):
            p_pm = row.cells[1].paragraphs[0]
            add_hyperlink(p_pm, pm_certs[i][0], pm_certs[i][1])

    # KPMGスタイルを適用
    set_table_style(cert_table)

    doc.add_paragraph()  # 適度な間隔を追加

    # 学習パス完了ステップ（統合版）
    add_heading_with_font(doc, '0.2 学習パス完了ステップ（完全版）', 2)

    p = doc.add_paragraph()
    p.add_run('予想所要時間: ').font.bold = True
    p.add_run('2～7日（Prerequisites有無により変動）')

    doc.add_paragraph()

    # Step 1
    heading = doc.add_paragraph()
    run1 = heading.add_run('Step 1: Prerequisites認証完了  ')
    run1.font.size = Pt(12)
    run1.font.bold = True
    set_run_font(run1)
    run2 = heading.add_run('【推奨・スキップ可】')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(192, 0, 0)  # 赤色
    set_run_font(run2)

    doc.add_paragraph('2つ以上の認証を完了（上記の推奨認証表を参照）', style='List Bullet')
    doc.add_paragraph('GitHub Foundations + Responsible AI（7～11時間）', style='List Bullet')
    doc.add_paragraph('Azure AI-900 + GitHub Foundations（10～16時間）', style='List Bullet')

    doc.add_paragraph()

    # Step 2
    heading = doc.add_paragraph()
    run1 = heading.add_run('Step 2: GitHub EMU + Pull Request提出  ')
    run1.font.size = Pt(12)
    run1.font.bold = True
    set_run_font(run1)
    run2 = heading.add_run('【必須】')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(192, 0, 0)  # 赤色
    set_run_font(run2)

    p = doc.add_paragraph(style='List Bullet')
    add_hyperlink(p, 'GitHub EMU',
                  'https://handbook.code.kpmg.com/KPMG-Code/GitHub/Organization%20onboarding/')
    p.add_run(' にオンボーディング完了')

    doc.add_paragraph('任意のリポジトリに最低1つのPull Requestを提出', style='List Bullet')
    doc.add_paragraph('所要時間: 1～2時間', style='List Bullet')

    doc.add_paragraph()

    # Step 3
    heading = doc.add_paragraph()
    run1 = heading.add_run('Step 3: Learning Path完了  ')
    run1.font.size = Pt(12)
    run1.font.bold = True
    set_run_font(run1)
    run2 = heading.add_run('【必須・二択一】')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(192, 0, 0)  # 赤色
    set_run_font(run2)

    # Learning Path比較テーブル
    doc.add_paragraph('以下のいずれかのLearning Pathを選択し、完了してください：')

    learning_table = doc.add_table(rows=5, cols=3)

    # ヘッダー行
    learning_headers = learning_table.rows[0].cells
    learning_headers[0].text = '項目'
    learning_headers[1].text = '開発者向け'
    learning_headers[2].text = 'PM向け'

    # Program ID行
    row1 = learning_table.rows[1].cells
    row1[0].text = 'Program ID'
    p_dev_id = row1[1].paragraphs[0]
    add_hyperlink(p_dev_id, 'GX25_CFS_DDF_AI_BLDG_WB_D_PRO',
                  'https://kpmgic.lms.hr.cloud.sap/learning/user/learning/program/viewProgramDetails.do?fromSF=Y&programID=GX25_CFS_DDF_AI_BLDG_WB_D_PRO')
    p_pm_id = row1[2].paragraphs[0]
    add_hyperlink(p_pm_id, 'GX25_CFS_DDF_AI_BLDG_WB_PM_PRO',
                  'https://kpmgic.lms.hr.cloud.sap/learning/user/learning/program/viewProgramDetails.do?fromSF=Y&programID=GX25_CFS_DDF_AI_BLDG_WB_PM_PRO')

    # 総時間行
    row2 = learning_table.rows[2].cells
    row2[0].text = '総時間'
    row2[1].text = '約5.3時間（318分）'
    row2[2].text = '約5.0時間（301分）'

    # モジュール内容行
    row3 = learning_table.rows[3].cells
    row3[0].text = 'モジュール内容'
    row3[1].text = '9モジュール:\n• Intro\n• AI Productivity\n• Inference API\n• Completion API\n• RAG（2部）\n• Feature Flags\n• Design\n• Resources'
    row3[2].text = '9モジュール:\n• Intro\n• Panel\n• AI Productivity\n• Why Workbench\n• IP Strategy\n• Microsoft Keynote\n• Migration\n• Feature Request\n• Support'

    # 備考行
    row4 = learning_table.rows[4].cells
    row4[0].text = '備考'
    row4[1].text = '重点: モジュール2・9の部署モード確認'
    row4[2].text = ''

    # KPMGスタイルを適用
    set_table_style(learning_table)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('⚠️ 重要: ').font.bold = True
    p.add_run('ビデオを95%以上視聴すること。完了ステータスは')
    run = p.add_run('24～48時間後')
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    p.add_run('にシステムに反映されます。')

    doc.add_paragraph()

    # Step 4
    heading = doc.add_paragraph()
    run1 = heading.add_run('Step 4: Knowledge Badge Assessment合格  ')
    run1.font.size = Pt(12)
    run1.font.bold = True
    set_run_font(run1)
    run2 = heading.add_run('【必須】')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(192, 0, 0)  # 赤色
    set_run_font(run2)

    doc.add_paragraph('全モジュール完了後、Assessment（試験）を受験', style='List Bullet')
    doc.add_paragraph('形式: 選択式問題（12問、30分）', style='List Bullet')
    doc.add_paragraph('合格基準: 70%以上', style='List Bullet')

    doc.add_paragraph()

    # Step 5
    heading = doc.add_paragraph()
    run1 = heading.add_run('Step 5: API Key取得  ')
    run1.font.size = Pt(12)
    run1.font.bold = True
    set_run_font(run1)
    run2 = heading.add_run('【必須】')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(192, 0, 0)  # 赤色
    set_run_font(run2)

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Assessment合格後、')
    add_hyperlink(p, 'Developer Onboarding Request Form',
                  'https://kpmggoprod.service-now.com/sp?id=sc_cat_item&sys_id=623c6518c314a61088532485e0013117&sysparm_category=3cae446893230a10324c76847aba1033')
    p.add_run(' からAPIキーをリクエスト')

    doc.add_paragraph('添付必須: Badge証明書 + Member Firm承認者のメール', style='List Bullet')
    doc.add_paragraph('発行期間: 2～3営業日', style='List Bullet')

    doc.add_paragraph()
    p_approver = doc.add_paragraph()
    run_approver = p_approver.add_run('⚠️ 日本地区のMember Firm Approver連絡先：')
    run_approver.font.bold = True
    run_approver.font.color.rgb = RGBColor(0, 112, 192)
    set_run_font(run_approver)

    # 日本地区Approverテーブル
    approver_table = doc.add_table(rows=4, cols=3)

    # ヘッダー行
    approver_headers = approver_table.rows[0].cells
    approver_headers[0].text = '氏名'
    approver_headers[1].text = '所属'
    approver_headers[2].text = 'メールアドレス'

    # Approverデータ
    japan_approvers = [
        ('Hotta Tomoyuki', 'FAS', 'Tomoyuki.Hotta@jp.kpmg.com'),
        ('Isoi Toru', 'AZSA', 'Toru.Isoi@jp.kpmg.com'),
        ('Tanaka Hidekazu', 'FAS', 'Hidekazu.Tanaka@jp.kpmg.com')
    ]

    for i, (name, dept, email) in enumerate(japan_approvers, 1):
        row = approver_table.rows[i].cells
        row[0].text = name
        row[1].text = dept
        row[2].text = email

    # KPMGスタイルを適用
    set_table_style(approver_table)

    doc.add_paragraph()  # 適度な間隔を追加

    # 時間コスト評価
    add_heading_with_font(doc, '0.3 時間コスト評価', 2)

    table = doc.add_table(rows=5, cols=3)

    headers = table.rows[0].cells
    headers[0].text = '学習モジュール'
    headers[1].text = '予定時間'
    headers[2].text = '実際の時間'

    modules = [
        ('事前要件（GitHub EMU等）', '___時間', ''),
        ('Developer / PM Learning Path', '5～5.3時間', ''),
        ('Assessment/Badge', '___時間', ''),
        ('合計', '___時間', '')
    ]

    for i, (module, est, actual) in enumerate(modules, 1):
        row = table.rows[i].cells
        row[0].text = module
        row[1].text = est
        row[2].text = actual

    # KPMGスタイルを適用
    set_table_style(table)

    doc.add_page_break()
