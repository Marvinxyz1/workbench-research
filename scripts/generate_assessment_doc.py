#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
KPMG Workbench 战略评估框架 Word文档生成器
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """为段落添加超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 设置超链接样式（蓝色下划线）
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

    return hyperlink

def set_heading_style(doc):
    """设置标题样式"""
    styles = doc.styles

    # 标题1样式
    h1 = styles['Heading 1']
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)

    # 标题2样式
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 112, 192)

    # 标题3样式
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True

def create_cover_page(doc):
    """创建封面页"""
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('KPMG Workbench 战略评估报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # 空行

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI开发平台深度评估框架')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 元信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('评估人员: [姓名]\n').font.size = Pt(12)
    info.add_run('职位: Senior Consultant, AI Development\n').font.size = Pt(12)
    info.add_run('日期: [评估日期]\n').font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 版本信息
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

def create_toc(doc):
    """创建目录页"""
    doc.add_heading('目录', 0)

    toc = doc.add_paragraph()
    toc.add_run('【请在Word中按以下步骤生成目录】\n').font.italic = True
    toc.add_run('1. 将光标放在此处\n').font.italic = True
    toc.add_run('2. 点击"引用"选项卡 → "目录" → 选择自动目录样式\n').font.italic = True
    toc.add_run('3. 完成内容填写后,右键点击目录 → "更新域" → "更新整个目录"\n\n').font.italic = True
    toc.add_run('注: 所有章节标题均已设置为标题样式,会自动生成超链接目录').font.italic = True

    doc.add_page_break()

def add_executive_summary(doc):
    """添加执行摘要"""
    doc.add_heading('执行摘要 (Executive Summary)', 1)

    doc.add_paragraph('【本节在完成所有评估后填写】')

    doc.add_heading('核心结论', 2)
    doc.add_paragraph('【3-5个要点总结】')

    doc.add_heading('关键发现', 2)
    doc.add_paragraph('✅ 主要优势:\n')
    doc.add_paragraph('⚠️ 主要限制:\n')
    doc.add_paragraph('🔥 核心亮点:\n')

    doc.add_heading('建议决策', 2)
    doc.add_paragraph('【Go/No-Go 建议 + 理由】')

    doc.add_page_break()

def add_dimension_0(doc):
    """维度0: 前置准备"""
    doc.add_heading('0. 前置准备: 学习路径与认证成本评估', 1)

    doc.add_paragraph('【评估背景】这是团队成员的最大入门障碍,需要详细评估时间成本和学习质量。')

    doc.add_heading('0.1 认证要求', 2)
    p = doc.add_paragraph('官方要求完成以下认证路径(来源: ')
    add_hyperlink(p, 'KPMG Workbench Learning & Development Hub',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-development.aspx')
    p.add_run('):')

    doc.add_paragraph('☐ Required Prerequisites/Certifications', style='List Bullet')
    doc.add_paragraph('☐ KPMG Workbench Knowledge Badge', style='List Bullet')
    doc.add_paragraph('☐ Developer Learning Path 或 Product Management Learning Path', style='List Bullet')

    doc.add_heading('0.2 时间成本评估', 2)

    # 创建表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '学习模块'
    headers[1].text = '预计时长'
    headers[2].text = '实际时长'

    modules = [
        ('Prerequisites', '', ''),
        ('Developer Learning Path', '', ''),
        ('Assessment/Badge', '', ''),
        ('总计', '', '')
    ]

    for i, (module, est, actual) in enumerate(modules, 1):
        row = table.rows[i].cells
        row[0].text = module
        row[1].text = est
        row[2].text = actual

    doc.add_paragraph()

    doc.add_heading('0.3 学习质量评估', 2)
    doc.add_paragraph('【高价值模块】\n')
    doc.add_paragraph('【可跳过模块】\n')
    doc.add_paragraph('【Tech Talks/Conference录像价值】\n')

    doc.add_heading('0.4 团队适配性预估', 2)
    doc.add_paragraph('【考虑"現状メンバーの空がない状態"(团队现在很忙)】\n')
    doc.add_paragraph('预计其他成员需要时长: ___小时\n')
    doc.add_paragraph('建议培训时间窗口: ___\n')
    doc.add_paragraph('技术背景较弱成员能否独立完成: ☐是 ☐否\n')

    doc.add_page_break()

def add_dimension_1(doc):
    """维度1: 技术能力"""
    doc.add_heading('1. 技术能力与效率评估', 1)

    doc.add_heading('1.1 技术栈兼容性', 2)

    doc.add_heading('设计系统 (Design Systems)', 3)
    p = doc.add_paragraph('参考: ')
    add_hyperlink(p, 'Design Systems for KPMG Workbench',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/Design-Systems-for-KPMG-Workbench.aspx')
    doc.add_paragraph('与现有UI框架兼容性: ___\n')
    doc.add_paragraph('迁移成本评估: ___\n')

    doc.add_heading('开发流程 (SDLC)', 3)
    p = doc.add_paragraph('参考: ')
    add_hyperlink(p, 'Software Development Lifecycle',
                  'https://docs.code.kpmg.com/GTK/Engineering-Ecosystem/Software-Development-Lifecycle-%28SDLC%29/sdlc/')
    doc.add_paragraph('与现有工作流冲突点: ___\n')
    doc.add_paragraph('CI/CD流程对比: ___\n')

    doc.add_heading('安全规范 (Secret Management)', 3)
    p = doc.add_paragraph('参考: ')
    add_hyperlink(p, 'Secret Management Best Practices',
                  'https://handbook.code.kpmg.com/digital-grc/secrets-management-best-practices/')
    doc.add_paragraph('相比现有方案优劣: ___\n')
    doc.add_paragraph('GitHub EMU授权流程复杂度: ___\n')

    doc.add_heading('1.2 易用性与学习曲线', 2)
    p = doc.add_paragraph('参考: ')
    add_hyperlink(p, 'KPMG Workbench User Guide',
                  'https://docs.code.kpmg.com/GTK/AI-Framework/KPMG-Workbench/')
    doc.add_paragraph('开发者体验(DX)评分: ___/10\n')
    doc.add_paragraph('文档完整度评分: ___/10\n')
    doc.add_paragraph('Hello World项目配置时间: ___小时\n')

    doc.add_heading('1.3 功能与限制', 2)
    doc.add_paragraph('【预装工具清单】\n')
    doc.add_paragraph('【相比现有环境的新功能】\n')
    doc.add_paragraph('【资源访问便捷性】\n')
    doc.add_paragraph('【配额限制与成本】\n')
    doc.add_paragraph('【最大短板】\n')

    doc.add_heading('1.4 开发效率提升(可量化)', 2)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '测试项目'
    headers[1].text = '现有流程'
    headers[2].text = 'Workbench'
    headers[3].text = '提升比例'

    items = [
        ('RAG Chatbot Demo', '', '', ''),
        ('从代码到部署', '', '', ''),
        ('Bug修复效率', '', '', '')
    ]

    for i, (item, old, new, imp) in enumerate(items, 1):
        row = table.rows[i].cells
        row[0].text = item
        row[1].text = old
        row[2].text = new
        row[3].text = imp

    doc.add_page_break()

def add_dimension_2(doc):
    """维度2: Agentic AI"""
    doc.add_heading('2. Agentic AI核心能力评估 🔥', 1)

    doc.add_paragraph('【重要性】这是Workbench的核心竞争力!')

    doc.add_heading('2.1 战略背景', 2)
    p = doc.add_paragraph('KPMG官方定位参考: ')
    add_hyperlink(p, 'The Agentic AI Advantage白皮书',
                  'https://kpmg.com/us/en/articles/2025/the-agentic-ai-advantage.html')

    doc.add_heading('2.2 Agent开发能力测试', 2)
    doc.add_paragraph('【是否支持Agent框架】\n')
    doc.add_paragraph('【预置Agent模板】\n')
    doc.add_paragraph('【相比LangChain/AutoGPT/CrewAI的优势】\n')

    doc.add_heading('2.3 真实场景测试 (必做!)', 2)
    doc.add_paragraph('测试任务: [例如:自动审计Agent]\n')
    doc.add_paragraph('开发时间: ___小时\n')
    doc.add_paragraph('Agent准确率: ___%\n')
    doc.add_paragraph('集成难度: ___\n')

    doc.add_heading('2.4 知识库集成', 2)
    doc.add_paragraph('【RAG能力】\n')
    doc.add_paragraph('【与KPMG知识库集成】\n')

    doc.add_heading('2.5 多Agent协作', 2)
    doc.add_paragraph('【是否支持多Agent】\n')
    doc.add_paragraph('【Orchestration机制】\n')

    doc.add_page_break()

def add_dimension_3(doc):
    """维度3: 商业价值"""
    doc.add_heading('3. 商业价值与客户应用', 1)

    doc.add_heading('3.1 Demo环境价值', 2)
    doc.add_paragraph('【售前加速】\n')
    doc.add_paragraph('Demo开发时间对比: [X周] → [Y天]\n')
    doc.add_paragraph('【提案竞争力提升】\n')
    doc.add_paragraph('【定制化能力】\n')

    doc.add_heading('3.2 内部服务开发', 2)
    doc.add_paragraph('【可开发的内部工具清单】\n')
    doc.add_paragraph('【与现有系统集成难度】\n')
    doc.add_paragraph('【知识沉淀机制】\n')

    doc.add_heading('3.3 客户项目边界与迁移成本 (风险!)', 2)
    doc.add_paragraph('【"不能用于客户服务"的边界】\n')
    doc.add_paragraph('Demo → 生产环境迁移成本: ___%\n')
    doc.add_paragraph('【数据隔离风险】\n')
    doc.add_paragraph('【合规性保证】\n')

    doc.add_page_break()

def add_dimension_4(doc):
    """维度4: 学习资源"""
    doc.add_heading('4. 学习资源与社区支持评估', 1)

    doc.add_heading('4.1 官方学习资源质量', 2)
    doc.add_paragraph('【Tech Talks系列 (2025年4-6月)】\n')
    doc.add_paragraph('最有价值的几期: ___\n')
    doc.add_paragraph('【Developers Conference录像 (2024年11月)】\n')
    doc.add_paragraph('【文档完整度】\n')

    doc.add_heading('4.2 社区与支持', 2)
    doc.add_paragraph('技术支持响应时间: ___小时\n')
    doc.add_paragraph('【Slack/Teams频道活跃度】\n')
    p = doc.add_paragraph('【Champions网络】参考: ')
    add_hyperlink(p, 'Global AI Ninjas and Navigators',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-AI/SitePages/Global-AI-Ninjas-Navigators.aspx')

    doc.add_heading('4.3 外部资源', 2)
    p = doc.add_paragraph('白皮书价值: ')
    add_hyperlink(p, 'AI Adoption in the Workplace',
                  'https://kpmg.com/au/en/insights/artificial-intelligence-ai/workplace-ai-adoption-success-insights-stories.html')
    doc.add_paragraph()
    p = doc.add_paragraph('Podcast价值: ')
    add_hyperlink(p, 'You Can with AI Podcast',
                  'https://kpmg.com/us/en/podcasts/you-can-with-ai.html')

    doc.add_page_break()

def add_dimension_5(doc):
    """维度5: 战略价值"""
    doc.add_heading('5. 战略价值与组织影响', 1)

    doc.add_heading('5.1 与KPMG AI愿景的契合度', 2)
    p = doc.add_paragraph('参考: ')
    add_hyperlink(p, 'KPMG Global AI战略',
                  'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-AI')
    doc.add_paragraph()
    p = doc.add_paragraph('Trusted AI原则: ')
    add_hyperlink(p, 'Trusted AI Principles',
                  'https://spo-global.kpmg.com/sites/go-oi-bus-People/SitePages/Trusted-AI.aspx')
    doc.add_paragraph('【AI Backbone定位价值】\n')
    doc.add_paragraph('【成为Champions案例的可能性】\n')

    doc.add_heading('5.2 跨团队协作机会', 2)
    doc.add_paragraph('【知识共享机制】\n')
    doc.add_paragraph('【资源复用可能性】\n')

    doc.add_heading('5.3 团队品牌与职业发展', 2)
    doc.add_paragraph('【内部可见性提升】\n')
    doc.add_paragraph('【个人成长机会】\n')

    doc.add_page_break()

def add_dimension_6(doc):
    """维度6: 风险评估"""
    doc.add_heading('6. 风险与合规性评估 ⚠️', 1)

    doc.add_heading('6.1 数据安全与隔离风险', 2)
    doc.add_paragraph('【数据泄露风险】\n')
    doc.add_paragraph('【多租户隔离机制】\n')
    doc.add_paragraph('【访问控制细粒度】\n')

    doc.add_heading('6.2 技术依赖与锁定风险', 2)
    doc.add_paragraph('【平台锁定风险】\n')
    doc.add_paragraph('迁移成本评估: ___\n')
    doc.add_paragraph('【技能转移性】\n')

    doc.add_heading('6.3 合规与审计', 2)
    doc.add_paragraph('【GDPR/数据主权合规性】\n')
    doc.add_paragraph('【审计日志完整性】\n')

    doc.add_heading('6.4 成本风险', 2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '成本项'
    headers[1].text = '预估金额'
    headers[2].text = '备注'

    items = [
        ('许可证费用', '', ''),
        ('算力成本(GPU)', '', ''),
        ('培训成本', '', ''),
        ('总计', '', '')
    ]

    for i, (item, cost, note) in enumerate(items, 1):
        row = table.rows[i].cells
        row[0].text = item
        row[1].text = cost
        row[2].text = note

    doc.add_paragraph()
    doc.add_paragraph('【机会成本】\n')

    doc.add_page_break()

def add_dimension_7(doc):
    """维度7: 推广运营"""
    doc.add_heading('7. 团队推广与运营考量', 1)

    doc.add_heading('7.1 团队适配性', 2)
    doc.add_paragraph('能快速上手的成员比例: ___%\n')
    doc.add_paragraph('【培训需求与成本】\n')
    doc.add_paragraph('【工作负载影响】\n')

    doc.add_heading('7.2 成本与收益(ROI)', 2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '指标'
    headers[1].text = '现状'
    headers[2].text = '使用Workbench后'

    items = [
        ('Demo开发时间', '', ''),
        ('项目中标率', '', ''),
        ('团队士气', '', ''),
        ('知识积累', '', '')
    ]

    for i, (metric, before, after) in enumerate(items, 1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = before
        row[2].text = after

    doc.add_paragraph()

    doc.add_heading('7.3 Go/No-Go决策标准 🎯', 2)

    doc.add_heading('全面推广条件 ✅', 3)
    doc.add_paragraph('☐ 单个成员完成认证时间 < 20小时', style='List Bullet')
    doc.add_paragraph('☐ Demo开发速度提升 > 30%', style='List Bullet')
    doc.add_paragraph('☐ 至少1个PoC获得客户正面反馈', style='List Bullet')
    doc.add_paragraph('☐ 技术支持响应时间 < 24小时', style='List Bullet')
    doc.add_paragraph('☐ 数据隔离通过安全审查', style='List Bullet')
    doc.add_paragraph('☐ 团队 > 70%认为值得学习', style='List Bullet')

    doc.add_heading('暂缓推广条件 ❌', 3)
    doc.add_paragraph('☐ 学习成本 > 1个月全职投入', style='List Bullet')
    doc.add_paragraph('☐ 隔离机制不清晰/有合规风险', style='List Bullet')
    doc.add_paragraph('☐ 团队 < 50%能独立完成开发', style='List Bullet')
    doc.add_paragraph('☐ 迁移成本 > 50%开发时间', style='List Bullet')
    doc.add_paragraph('☐ 官方支持不足', style='List Bullet')
    doc.add_paragraph('☐ 有更好的替代方案', style='List Bullet')

    doc.add_heading('7.4 试点方案 (Phased Rollout)', 2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '阶段'
    headers[1].text = '时长'
    headers[2].text = '参与人员'
    headers[3].text = '交付物'

    phases = [
        ('Phase 1: 个人探索', '1个月', 'SC级别1人', '技术评估报告'),
        ('Phase 2: 小组试点', '2个月', 'SC + 1-2成员', 'PoC + ROI数据'),
        ('Phase 3: 决策点', '1周', '管理层', 'Go/No-Go决策'),
        ('Phase 4: 全面推广', '3个月', '全团队', '培训完成+持续优化')
    ]

    for i, (phase, duration, people, output) in enumerate(phases, 1):
        row = table.rows[i].cells
        row[0].text = phase
        row[1].text = duration
        row[2].text = people
        row[3].text = output

    doc.add_page_break()

def add_appendix_references(doc):
    """添加参考资源附录"""
    doc.add_heading('附录: 参考资源', 1)

    doc.add_heading('A. 官方学习资源', 2)

    resources = [
        ('KPMG Workbench Learning & Development Hub',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-development.aspx'),
        ('Developer Learning Path',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-development-track.aspx'),
        ('Product Management Learning Path',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/KPMG-Workbench-learning-and-development-product-management-track.aspx'),
    ]

    for title, url in resources:
        p = doc.add_paragraph('• ', style='List Bullet')
        add_hyperlink(p, title, url)

    doc.add_heading('B. 技术文档', 2)

    tech_docs = [
        ('KPMG Workbench User Guide',
         'https://docs.code.kpmg.com/GTK/AI-Framework/KPMG-Workbench/'),
        ('Design Systems for KPMG Workbench',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/SitePages/Design-Systems-for-KPMG-Workbench.aspx'),
        ('Software Development Lifecycle (SDLC)',
         'https://docs.code.kpmg.com/GTK/Engineering-Ecosystem/Software-Development-Lifecycle-%28SDLC%29/sdlc/'),
        ('Secret Management Best Practices',
         'https://handbook.code.kpmg.com/digital-grc/secrets-management-best-practices/'),
    ]

    for title, url in tech_docs:
        p = doc.add_paragraph('• ', style='List Bullet')
        add_hyperlink(p, title, url)

    doc.add_heading('C. 战略资源', 2)

    strategic = [
        ('KPMG Global aIQ Hub',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-AI'),
        ('Global AI Ninjas and Navigators',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-AI/SitePages/Global-AI-Ninjas-Navigators.aspx'),
        ('Trusted AI Learning Path',
         'https://spo-global.kpmg.com/sites/go-oi-bus-People/SitePages/Trusted-AI.aspx'),
    ]

    for title, url in strategic:
        p = doc.add_paragraph('• ', style='List Bullet')
        add_hyperlink(p, title, url)

    doc.add_heading('D. 白皮书与洞察', 2)

    insights = [
        ('The Agentic AI Advantage',
         'https://kpmg.com/us/en/articles/2025/the-agentic-ai-advantage.html'),
        ('AI Adoption in the Workplace',
         'https://kpmg.com/au/en/insights/artificial-intelligence-ai/workplace-ai-adoption-success-insights-stories.html'),
        ('KPMG Revolutionizes AI Delivery',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-AI/SitePages/KPMG-revolutionizes-AI-delivery-with-a-first-of-its-kind-global-AI-platform.aspx'),
    ]

    for title, url in insights:
        p = doc.add_paragraph('• ', style='List Bullet')
        add_hyperlink(p, title, url)

    doc.add_heading('E. 会议录像 (推荐)', 2)

    videos = [
        ('Microsoft Keynote - Agentic AI Thinking',
         'https://spo-global.kpmg.com/:v:/r/sites/GO-OI-BUS-GTK-WB/KPMGWorkbenchDevCon/Microsoft%20Keynote%20recording%20-%20Agentic%20AI%20Thinking.mp4'),
        ('KPMG Keynote - Workbench Champions',
         'https://spo-global.kpmg.com/sites/GO-OI-BUS-GTK-WB/KPMGWorkbenchDevCon/KPMG%20Keynote%20recording%20-%20Workbench%20Champions.mp4'),
        ('Q&A with Product Owners',
         'https://spo-global.kpmg.com/:v:/r/sites/GO-OI-BUS-GTK-WB/KPMGWorkbenchDevCon/The%20Open%20Forum%20Live%20Q%26A%20with%20Workbench%20Product.mp4'),
    ]

    for title, url in videos:
        p = doc.add_paragraph('• ', style='List Bullet')
        add_hyperlink(p, title, url)

    doc.add_heading('F. 外部资源', 2)

    p = doc.add_paragraph('• ', style='List Bullet')
    add_hyperlink(p, 'You Can with AI Podcast',
                  'https://kpmg.com/us/en/podcasts/you-can-with-ai.html')

    doc.add_page_break()

def add_appendix_timeline(doc):
    """添加评估时间表"""
    doc.add_heading('附录: 评估时间表 (建议)', 1)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = table.rows[0].cells
    headers[0].text = '周次'
    headers[1].text = '任务'
    headers[2].text = '交付物'

    timeline = [
        ('W1', '完成Prerequisites + Developer Learning Path', '认证徽章'),
        ('W2', '开发测试Demo 1 (如: RAG Chatbot)', '技术能力评估数据'),
        ('W3', '开发测试Demo 2 (如: Audit Agent)', 'Agentic AI能力评估'),
        ('W4', '风险分析、ROI计算、撰写报告', '完整评估报告 + Executive Summary')
    ]

    for i, (week, task, output) in enumerate(timeline, 1):
        row = table.rows[i].cells
        row[0].text = week
        row[1].text = task
        row[2].text = output

def main():
    """主函数"""
    import sys
    import io

    # 修复Windows控制台编码问题
    if sys.platform == 'win32':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    print('开始生成KPMG Workbench战略评估报告...')

    # 创建文档
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 设置标题样式
    set_heading_style(doc)

    # 创建各个部分
    print('  - 创建封面页...')
    create_cover_page(doc)

    print('  - 创建目录页...')
    create_toc(doc)

    print('  - 添加执行摘要...')
    add_executive_summary(doc)

    print('  - 添加维度0: 前置准备...')
    add_dimension_0(doc)

    print('  - 添加维度1: 技术能力...')
    add_dimension_1(doc)

    print('  - 添加维度2: Agentic AI...')
    add_dimension_2(doc)

    print('  - 添加维度3: 商业价值...')
    add_dimension_3(doc)

    print('  - 添加维度4: 学习资源...')
    add_dimension_4(doc)

    print('  - 添加维度5: 战略价值...')
    add_dimension_5(doc)

    print('  - 添加维度6: 风险评估...')
    add_dimension_6(doc)

    print('  - 添加维度7: 推广运营...')
    add_dimension_7(doc)

    print('  - 添加参考资源附录...')
    add_appendix_references(doc)

    print('  - 添加评估时间表...')
    add_appendix_timeline(doc)

    # 保存文档
    output_path = r'C:\Users\junchenma\workbench-research\KPMG_Workbench战略评估报告.docx'
    doc.save(output_path)

    print(f'✅ 文档生成成功!')
    print(f'📁 文件位置: {output_path}')
    print('\n📝 后续步骤:')
    print('1. 在Word中打开文档')
    print('2. 将光标放在"目录"页的占位符处')
    print('3. 点击 引用 → 目录 → 选择自动目录样式')
    print('4. 完成内容填写后,右键目录 → 更新域 → 更新整个目录')
    print('5. 使用Word的"另存为PDF"功能导出最终版本')

if __name__ == '__main__':
    main()
