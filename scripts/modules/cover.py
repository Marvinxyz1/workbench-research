# -*- coding: utf-8 -*-
"""
KPMG Workbench 戦略評価フレームワーク - 表紙・目次モジュール
"""

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_cover_page(doc):
    """表紙ページを作成

    Args:
        doc: python-docx文書オブジェクト
    """
    # タイトル
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('KPMG Workbench 戦略評価レポート')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()  # 空行

    # サブタイトル
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI開発プラットフォーム詳細評価フレームワーク')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # メタ情報
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('評価担当者: [氏名]\n').font.size = Pt(12)
    info.add_run('役職: Senior Consultant, AI Development\n').font.size = Pt(12)
    info.add_run('日付: [評価日]\n').font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # バージョン情報
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def create_toc(doc):
    """目次ページを作成

    Args:
        doc: python-docx文書オブジェクト
    """
    doc.add_heading('目次', 0)

    toc = doc.add_paragraph()
    toc.add_run('【Wordで目次を生成する手順】\n').font.italic = True
    toc.add_run('1. カーソルをここに置く\n').font.italic = True
    toc.add_run('2. 「参考資料」タブ → 「目次」 → 自動目次スタイルを選択\n').font.italic = True
    toc.add_run('3. 内容記入完了後、目次を右クリック → 「フィールド更新」 → 「目次をすべて更新」\n\n').font.italic = True
    toc.add_run('注: すべての章見出しは見出しスタイルに設定されており、自動的にハイパーリンク目次が生成されます').font.italic = True

    doc.add_page_break()


def add_executive_summary(doc):
    """エグゼクティブサマリーを追加

    Args:
        doc: python-docx文書オブジェクト
    """
    doc.add_heading('エグゼクティブサマリー (Executive Summary)', 1)

    doc.add_paragraph('【本セクションはすべての評価完了後に記入】')

    doc.add_heading('核心結論', 2)
    doc.add_paragraph('【3～5つの要点まとめ】')

    doc.add_heading('主要発見事項', 2)
    doc.add_paragraph('✅ 主な強み:\n')
    doc.add_paragraph('⚠️ 主な制限:\n')
    doc.add_paragraph('🔥 コアハイライト:\n')

    doc.add_heading('推奨決定', 2)
    doc.add_paragraph('【Go/No-Go 推奨 + 理由】')

    doc.add_page_break()
